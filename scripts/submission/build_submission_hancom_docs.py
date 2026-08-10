from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor
from docx.table import Table


GENERATED_AT_UTC = datetime.now(timezone.utc).replace(microsecond=0)
ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "output" / "docx"
SCREENSHOT = ROOT / "docs" / "evidence" / "screenshots" / "title-screen-mwo-meogeul-geonyang-320x568.png"
GAMEPLAY = ROOT / "docs" / "evidence" / "food-visual-slice-gameplay.png"
LOBBY = ROOT / "docs" / "evidence" / "screenshots" / "bright-room-lobby-412x915.png"
RESULTS = ROOT / "docs" / "evidence" / "screenshots" / "bright-room-results-412x915.png"

GAME_DOCX = OUTPUT_DIR / "뭐_먹을_거냥_게임_소개_및_설명_편집원본.docx"
AI_DOCX = OUTPUT_DIR / "뭐_먹을_거냥_AI_활용_기술_문서_편집원본.docx"

# standard_business_brief preset with explicit Korean/A4 named overrides.
FONT = "맑은 고딕"
CODE_FONT = "Consolas"
INK = "241B35"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
PINK = "D6336C"
MUTED = "66616F"
LINE = "D7D3DE"
LIGHT_GRAY = "F2F4F7"
LIGHT_PINK = "FFF0F5"
LILAC = "F3F0FF"
MINT = "EAF8F1"
CREAM = "FFF8E7"
YELLOW = "FFF3BF"
WHITE = "FFFFFF"
PLACEHOLDER_FILL = "FFF2CC"

PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
MARGIN_LEFT_MM = 16
MARGIN_RIGHT_MM = 16
MARGIN_TOP_MM = 15
MARGIN_BOTTOM_MM = 18
CONTENT_WIDTH_MM = PAGE_WIDTH_MM - MARGIN_LEFT_MM - MARGIN_RIGHT_MM
CONTENT_WIDTH_DXA = round(CONTENT_WIDTH_MM / 25.4 * 1440)
TABLE_INDENT_DXA = 120


def set_run_font(run, *, name: str = FONT, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_borders(cell, color: str = LINE, size: int = 5) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_mm: list[float]) -> None:
    if round(sum(widths_mm), 3) != round(CONTENT_WIDTH_MM, 3):
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_MM} mm: {widths_mm}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    dxa_widths = [round(width / 25.4 * 1440) for width in widths_mm]
    dxa_widths[-1] += CONTENT_WIDTH_DXA - sum(dxa_widths)
    for width in dxa_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        for cell, width in zip(row.cells, dxa_widths, strict=True):
            cell.width = Mm(width / 1440 * 25.4)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_borders(cell)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(5)
    pf.line_spacing = 1.15

    for name, size, color, before, after in (
        ("Title", 27, INK, 0, 7),
        ("Subtitle", 14, MUTED, 0, 8),
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 10, 5),
        ("Heading 3", 11.5, DARK_BLUE, 7, 4),
    ):
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = FONT
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.paragraph_format.left_indent = Inches(0.50)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def configure_page(doc: Document, *, title: str, subject: str) -> None:
    section = doc.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.left_margin = Mm(MARGIN_LEFT_MM)
    section.right_margin = Mm(MARGIN_RIGHT_MM)
    section.top_margin = Mm(MARGIN_TOP_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)
    section.different_first_page_header_footer = True

    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "뭐 먹을 거냥? 개인 참가자"
    props.keywords = "NHN 게임 제작 해커톤, 개인 참가, 뭐 먹을 거냥?"
    props.comments = "한컴 한글 편집 원본. 최종 링크와 제출자 정보를 입력한 뒤 PDF로 내보내세요."
    props.created = GENERATED_AT_UTC
    props.modified = GENERATED_AT_UTC


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def add_running_header_footer(doc: Document, label: str) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label)
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("뭐 먹을 거냥?  ·  ")
    set_run_font(run, size=8.5, color=MUTED)
    add_field(p, "PAGE")


def add_bottom_border(paragraph, color: str = LINE, size: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_cover(doc: Document, *, subtitle: str, status: str, lead: str, include_image: bool = True) -> None:
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_before = Pt(14)
    kicker.paragraph_format.space_after = Pt(8)
    run = kicker.add_run("NHN 게임 제작 해커톤  ·  개인 참가")
    set_run_font(run, size=10, color=PINK, bold=True)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("뭐 먹을 거냥?")
    set_run_font(run, size=27, color=INK, bold=True)

    sub = doc.add_paragraph(style="Subtitle")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    set_run_font(run, size=14, color=MUTED, bold=True)

    status_p = doc.add_paragraph()
    status_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    status_p.paragraph_format.space_before = Pt(3)
    status_p.paragraph_format.space_after = Pt(12)
    add_paragraph_shading(status_p, YELLOW)
    run = status_p.add_run(status)
    set_run_font(run, size=9.5, color=INK, bold=True)

    if include_image and SCREENSHOT.exists():
        image_p = doc.add_paragraph()
        image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_p.paragraph_format.space_after = Pt(8)
        shape = image_p.add_run().add_picture(str(SCREENSHOT), width=Mm(49))
        shape._inline.docPr.set("descr", "뭐 먹을 거냥? 모바일 타이틀 화면")
        shape._inline.docPr.set("name", "게임 타이틀 화면")

    lead_p = doc.add_paragraph()
    lead_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead_p.paragraph_format.left_indent = Mm(15)
    lead_p.paragraph_format.right_indent = Mm(15)
    lead_p.paragraph_format.space_before = Pt(3)
    lead_p.paragraph_format.space_after = Pt(8)
    add_paragraph_shading(lead_p, LIGHT_PINK)
    run = lead_p.add_run(lead)
    set_run_font(run, size=10.5, color=INK, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(8)
    run = meta.add_run("한컴 편집용 DOCX 원본  ·  최종 링크 입력 후 PDF로 내보내기")
    set_run_font(run, size=8.5, color=MUTED)


def add_page_title(doc: Document, number: str, title: str) -> None:
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"{number}  {title}")
    set_run_font(run, size=16, color=BLUE, bold=True)
    add_bottom_border(p, color=LINE, size=8)


def add_body(doc: Document, text: str, *, bold: bool = False, color: str | None = None,
             align=WD_ALIGN_PARAGRAPH.LEFT, after: float = 5) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=10.5, color=color or INK, bold=bold)


def add_subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=13, color=BLUE, bold=True)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.50)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3.5)
        run = p.add_run(item)
        set_run_font(run, size=10.2, color=INK)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.50)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        set_run_font(run, size=10.2, color=INK)


def add_note(doc: Document, title: str, text: str, *, fill: str = LIGHT_PINK) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Mm(3)
    p.paragraph_format.right_indent = Mm(3)
    add_paragraph_shading(p, fill)
    r = p.add_run(f"{title}  ")
    set_run_font(r, size=10.2, color=INK, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attr}"), FONT)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    r_pr.extend([fonts, color, underline, size])
    run.append(r_pr)
    node = OxmlElement("w:t")
    node.text = text
    run.append(node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def fill_cell(cell, text: str, *, bold: bool = False, size: float = 9.5,
              color: str = INK, align=WD_ALIGN_PARAGRAPH.LEFT, placeholder: bool = False) -> None:
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)
    if placeholder:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), PLACEHOLDER_FILL)
        run._element.get_or_add_rPr().append(shd)


def add_table(doc: Document, rows: list[list[str]], widths_mm: list[float], *, header: bool = True,
              font_size: float = 9.2, first_col_fill: str | None = None,
              placeholder_cells: set[tuple[int, int]] | None = None) -> Table:
    table = doc.add_table(rows=len(rows), cols=len(widths_mm))
    placeholder_cells = placeholder_cells or set()
    for r_idx, row_data in enumerate(rows):
        if len(row_data) != len(widths_mm):
            raise ValueError(f"Row width mismatch: {row_data}")
        for c_idx, text in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            if header and r_idx == 0:
                set_cell_shading(cell, INK)
                fill_cell(cell, text, bold=True, size=font_size, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            else:
                if first_col_fill and c_idx == 0:
                    set_cell_shading(cell, first_col_fill)
                fill_cell(cell, text, bold=bool(first_col_fill and c_idx == 0), size=font_size,
                          placeholder=(r_idx, c_idx) in placeholder_cells)
    if header:
        set_repeat_table_header(table.rows[0])
    set_table_geometry(table, widths_mm)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_link_table(doc: Document, *, include_youtube: bool) -> None:
    rows = [
        ("소스 코드", "https://github.com/h200082/oneul-mwo-sseol", "https://github.com/h200082/oneul-mwo-sseol"),
        ("플레이 링크", "https://h200082.github.io/oneul-mwo-sseol/", "https://h200082.github.io/oneul-mwo-sseol/"),
    ]
    if include_youtube:
        rows.append(("플레이 영상", "[YouTube 최종 URL 입력]", None))
    table = doc.add_table(rows=len(rows), cols=2)
    for index, (label, visible, url) in enumerate(rows):
        set_cell_shading(table.cell(index, 0), LILAC)
        fill_cell(table.cell(index, 0), label, bold=True, size=9.6)
        p = table.cell(index, 1).paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if url:
            add_hyperlink(p, visible, url)
            if label == "플레이 링크":
                r = p.add_run("  (실제 배포·시크릿 창·모바일 접속 확인 전)")
                set_run_font(r, size=8.5, color="C14A2A", bold=True)
        else:
            r = p.add_run(visible)
            set_run_font(r, size=9.6, color="C14A2A", bold=True)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), PLACEHOLDER_FILL)
            r._element.get_or_add_rPr().append(shd)
    set_table_geometry(table, [35, 143])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_picture(doc: Document, path: Path, width_mm: float, alt_text: str, caption: str | None = None) -> None:
    if not path.exists():
        add_note(doc, "이미지 누락", str(path), fill=YELLOW)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    shape = p.add_run().add_picture(str(path), width=Mm(width_mm))
    shape._inline.docPr.set("descr", alt_text)
    shape._inline.docPr.set("name", alt_text)
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(5)
        run = cp.add_run(caption)
        set_run_font(run, size=8.5, color=MUTED, italic=True)


def add_picture_pair(doc: Document, left: Path, right: Path, *, width_mm: float,
                     left_alt: str, right_alt: str, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    left_shape = p.add_run().add_picture(str(left), width=Mm(width_mm))
    left_shape._inline.docPr.set("descr", left_alt)
    left_shape._inline.docPr.set("name", left_alt)
    p.add_run("      ")
    right_shape = p.add_run().add_picture(str(right), width=Mm(width_mm))
    right_shape._inline.docPr.set("descr", right_alt)
    right_shape._inline.docPr.set("name", right_alt)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(4)
    run = cp.add_run(caption)
    set_run_font(run, size=8.5, color=MUTED, italic=True)


def add_final_information_page(doc: Document, *, include_youtube: bool) -> None:
    doc.add_page_break()
    add_page_title(doc, "FINAL", "최종 제출 정보 입력")
    add_note(doc, "편집 안내", "노란색 대괄호 항목을 최종 값으로 교체한 뒤, 한컴 한글에서 PDF로 내보내고 모든 링크를 다시 확인하세요.", fill=YELLOW)
    rows = [
        ["항목", "최종 입력값"],
        ["참가 형태", "개인 참가 (팀원 롤 기술서 제출 생략)"],
        ["제출자 이름", "[제출자 이름 입력]"],
        ["GitHub 저장소", "https://github.com/h200082/oneul-mwo-sseol"],
        ["플레이 URL", "[검증된 최종 Pages URL 입력]"],
    ]
    if include_youtube:
        rows.append(["YouTube URL", "[30~60초 플레이 영상 URL 입력]"])
    rows.extend([
        ["최종 commit SHA", "[최종 commit SHA 입력]"],
        ["문서 생성일", "2026-08-10"],
        ["실제 제출일·시각", "[제출일·시각 입력]"],
        ["최종 메모", "[메모 입력]"],
    ])
    placeholder_rows = {(i, 1) for i, row in enumerate(rows) if row[1].startswith("[")}
    add_table(doc, rows, [42, 136], header=True, font_size=9.3, first_col_fill=LILAC,
              placeholder_cells=placeholder_rows)
    add_subheading(doc, "최종 확인")
    checks = [
        "[ ] 시크릿 창과 PC·실제 모바일에서 플레이 URL을 열었다.",
        "[ ] GitHub 저장소와 제출 commit SHA가 일치한다.",
        "[ ] 문서 속 모든 링크를 클릭해 확인했다.",
        "[ ] 한컴에서 DOCX를 저장한 뒤 다시 열어 글자·표·이미지를 확인했다.",
        "[ ] 한컴 저장 후 URL이 평문이면 하이퍼링크를 다시 지정했다.",
        "[ ] 최종 PDF를 전 페이지 검토했다.",
    ]
    if include_youtube:
        checks.insert(2, "[ ] YouTube 영상이 공개 또는 일부 공개이며 30~60초다.")
    add_bullets(doc, checks)
    add_note(doc, "5번 제출물", "개인 참가이므로 팀원 롤 기술서는 해당 없음으로 생략합니다.", fill=MINT)


def base_document(*, title: str, subject: str, label: str) -> Document:
    doc = Document()
    configure_styles(doc)
    configure_page(doc, title=title, subject=subject)
    add_running_header_footer(doc, label)
    return doc


def build_game_document() -> Path:
    doc = base_document(
        title="뭐 먹을 거냥? - 게임 소개 및 설명 문서",
        subject="NHN 게임 제작 해커톤 게임 소개 및 설명 문서",
        label="뭐 먹을 거냥? · 게임 소개 및 설명",
    )
    add_cover(
        doc,
        subtitle="게임 소개 및 설명 문서",
        status="사전 제출본 · GitHub Pages와 YouTube 링크 최종 확인 전",
        lead="먹고 싶은 메뉴는 포획하고 나머지는 정확히 반으로 썰어, 최대 8명의 식사 내기 순위와 메뉴 취향을 함께 확인하는 모바일 우선 브라우저 파티게임.",
    )

    doc.add_page_break()
    add_page_title(doc, "01", "게임의 목표와 핵심 재미")
    add_body(doc, "한 판에는 전체 음식 50종 중 중복 없는 20종이 등장한다. 플레이어는 나머지 음식을 실제 그림의 실루엣 기준으로 최대한 정확히 반으로 가르고, 먹고 싶은 음식은 최대 2개까지 길게 눌러 포획한다.")
    add_table(doc, [
        ["항목", "내용"],
        ["플랫폼", "모바일·PC 웹 브라우저"],
        ["플레이 인원", "혼자 하기 또는 온라인 방 2~8명"],
        ["한 판 길이", "약 45~60초"],
        ["출제", "전체 50종 중 중복 없는 20종"],
    ], [40, 138], header=True, first_col_fill=LILAC)
    add_bullets(doc, [
        "베는 손맛: 음식 알파 실루엣의 양쪽 면적을 직접 비교해 50:50에 가까울수록 높은 점수를 얻는다.",
        "메뉴 선택: 먹고 싶은 음식은 점수 경쟁에서 제외하고 최대 2개까지 내 취향으로 포획한다.",
        "짧은 파티: 같은 20종으로 최대 8명이 경쟁하고 순위·겹친 취향·고정 내기를 함께 확인한다.",
    ])
    add_subheading(doc, "점수 계산")
    add_table(doc, [["베기 정확도", "100 × (1 - |A - B| ÷ (A + B))"]], [46, 132], header=False, first_col_fill=LIGHT_PINK)
    add_bullets(doc, [
        "A와 B는 베기 선 양쪽에 남은 실제 음식의 가중 알파 면적이다.",
        "놓친 음식은 0점이며, 포획한 음식은 평균 정확도 계산에서 제외된다.",
        "포획 0·1·2개일 때 평균의 분모는 각각 20·19·18이다.",
        "20개를 베기·포획·놓침 중 하나로 모두 처리하면 라운드가 끝난다.",
    ])

    doc.add_page_break()
    add_page_title(doc, "02", "조작과 라운드 진행")
    add_numbers(doc, [
        "드래그해서 베기: 음식의 한쪽 테두리에서 반대쪽 테두리까지 드래그한다. 입력 중에도 음식은 계속 낙하한다.",
        "0.32초 길게 눌러 포획: 먹고 싶은 음식 위에서 움직이지 않고 누른다. 한 판에 최대 2개이며 사용하지 않아도 된다.",
        "움직이면 즉시 베기로 전환: 누르는 중 시작점에서 14px 이상 움직이면 포획 대기를 취소하고 같은 입력을 베기로 해석한다.",
        "짧은 탭은 무효: 짧은 탭과 취소 입력은 라운드 진행이나 포획 슬롯을 소비하지 않는다.",
    ])
    add_picture(doc, GAMEPLAY, 45, "실제 음식 알파 실루엣을 베는 플레이 화면", "실제 음식 실루엣 기준의 베기 판정")
    add_subheading(doc, "점점 빨라지는 20개")
    add_table(doc, [
        ["구간", "낙하 시간", "움직임"],
        ["1~5", "2.6초", "회전 없음"],
        ["6~15", "2.2초", "회전 증가"],
        ["16~20", "1.8초", "회전 + 마지막 2개 좌우 이동"],
    ], [32, 40, 106], header=True)
    add_note(doc, "진행 안내", "16번째 음식 전에는 FINAL 5 안내가 나타난다. 정확도 등급은 95점 이상 '칼각', 80점 이상 '훌륭', 60점 이상 '좋아', 그 아래 '아쉬워'로 표시된다.", fill=CREAM)

    doc.add_page_break()
    add_page_title(doc, "03", "혼자 또는 2~8명이 함께")
    add_picture_pair(doc, LOBBY, RESULTS, width_mm=42, left_alt="방 코드와 QR이 보이는 대기실", right_alt="순위와 포획 메뉴가 보이는 공동 결과", caption="왼쪽: 대기실  ·  오른쪽: 공동 결과")
    add_bullets(doc, [
        "방 생성 후 QR 코드, 초대 링크 또는 8자리 코드로 참가한다.",
        "준비 버튼 없이 방장이 시작하면 참가자 명단을 잠그고 같은 20종 덱을 선로딩한다.",
        "자동 선로딩 뒤 공통 4초 카운트다운으로 시작한다.",
        "최고 평균 정확도가 1등, 최저 평균 정확도가 꼴찌이며 동점은 공동 순위다.",
        "결과에서 개인 포획 메뉴, 1등의 포획 메뉴, 정확히 겹친 메뉴 또는 가까운 카테고리 취향을 확인한다.",
        "고정 내기: 꼴찌가 1등의 식사 1인분을 부담한다. 공동 순위일 때는 참가자가 부담 방식을 합의한다.",
        "게임은 최종 식사 메뉴를 자동으로 강제하지 않는다.",
    ])

    doc.add_page_break()
    add_page_title(doc, "04", "실행 방법과 제출 링크")
    add_subheading(doc, "브라우저에서 바로 플레이")
    add_bullets(doc, [
        "최신 Chrome, Edge, Safari 등에서 플레이 링크를 연다.",
        "소리·나레이션·진동은 기본 ON이며 설정에서 끌 수 있다.",
        "QR 카메라 스캔은 HTTPS에서 사용한다. 링크·8자리 코드 참가는 카메라 없이 가능하다.",
    ])
    add_subheading(doc, "소스에서 실행")
    add_table(doc, [
        ["환경", "명령"],
        ["Node.js 20.19 이상", "npm ci\nnpm run dev\n\nnpm run build\nnpm run preview"],
    ], [50, 128], header=True, font_size=9.5)
    add_subheading(doc, "제출 링크")
    add_link_table(doc, include_youtube=True)
    add_note(doc, "최종 제출 전 교체·확인", "GitHub Pages workflow 성공 → 시크릿 창·PC·실제 모바일 접속 → 서로 다른 2개 브라우저/기기의 방 플레이 → YouTube 링크 입력 → PDF 다시 생성 → 모든 링크 클릭 검사.", fill=LIGHT_PINK)
    add_body(doc, "개인 참가이므로 팀원 롤 기술서는 제출 대상에서 제외한다.", bold=True, color=DARK_BLUE)

    add_final_information_page(doc, include_youtube=True)
    doc.save(GAME_DOCX)
    return GAME_DOCX


def build_ai_document() -> Path:
    doc = base_document(
        title="뭐 먹을 거냥? - AI 활용 기술 문서",
        subject="NHN 게임 제작 해커톤 AI 활용 기술 문서",
        label="뭐 먹을 거냥? · AI 활용 기술",
    )
    add_cover(
        doc,
        subtitle="AI 활용 기술 문서",
        status="사전 제출본 · 최종 배포 링크와 commit SHA 확정 전",
        lead="제작 단계에서는 적극적으로, 실행 중에는 호출하지 않게. AI를 기획·구현·이미지·음성 제작과 검증에 사용하되, 규칙과 최종 자산은 사람이 결정했다.",
    )

    doc.add_page_break()
    add_page_title(doc, "01", "도구와 적용 범위")
    add_note(doc, "핵심 원칙", "게임 런타임에는 생성형 AI 호출, Azure 키·엔드포인트·Speech SDK가 없다. AI 결과는 사람이 검수한 정적 이미지와 정적 음성으로만 포함한다.", fill=MINT)
    add_table(doc, [
        ["도구", "활용", "최종 반영"],
        ["OpenAI Codex", "기획 구조화, TypeScript 구현, 테스트, 디버깅, 문서·프롬프트 기록", "사람이 diff·실행 결과·회귀를 검토한 소스만 채택"],
        ["Codex 내장 ImageGen", "음식 50종과 타이틀 고양이 셰프 시안", "사람이 선별하고 로컬 후처리한 투명 WebP"],
        ["Azure AI Speech", "한국어 음식 나레이션 후보", "A/B·블라인드 청취 후 정적 MP3/WAV 50개"],
        ["Pillow 12.3.0", "크로마키 제거, despill, 알파 크롭, 리사이즈, WebP 검사", "로컬 재현 가능한 스크립트"],
    ], [34, 72, 72], header=True, font_size=8.7)
    add_subheading(doc, "개발 구조")
    add_table(doc, [
        ["계층", "역할"],
        ["DOM", "AppController - 화면·로비·결과 UI"],
        ["게임", "Phaser PrototypeScene - 낙하·베기·포획"],
        ["도메인", "점수·기하·방·결과 순수 로직"],
        ["게이트웨이", "localStorage·BroadcastChannel / Firebase Auth·Firestore"],
    ], [42, 136], header=True, first_col_fill=LILAC, font_size=9.1)
    add_bullets(doc, [
        "QR은 qrcode로 생성하고 브라우저 BarcodeDetector 지원 시 카메라 스캔을 제공한다.",
        "Web Audio·Vibration API·정적 나레이션이 감각 피드백을 담당한다.",
        "TypeScript·Vite 빌드를 GitHub Pages 배포 artifact로 만든다.",
    ])

    doc.add_page_break()
    add_page_title(doc, "02", "Codex 개발 협업과 프롬프트")
    add_body(doc, "Codex에는 결과만 요청하지 않고 목표, 금지 조건, 변경 범위와 검증 기준을 함께 제공했다. 제안된 구현은 사람의 선택과 자동 테스트를 모두 통과한 경우에만 반영했다.")
    add_subheading(doc, "대표 프롬프트와 사람의 결정")
    add_table(doc, [
        ["주제", "주요 지시", "사람의 최종 판단"],
        ["게임 규칙", "게임 규칙은 프레임워크에서 분리하고 경계값을 자동 테스트한다.", "포획 제외 평균, 놓침 0점, 순위·내기 규칙을 별도 도메인으로 고정"],
        ["조작 충돌", "드래그 중 음식이 멈추지 않게 하고 베기와 포획 입력이 충돌하지 않도록 한다.", "더블클릭·일시정지를 제외하고 0.32초 hold + 14px 이동 전환 채택"],
        ["실루엣 점수", "원형이 아닌 음식 그림 그대로 판정하고 그 차이를 점수로 전달한다.", "128×128 가중 알파 마스크, 실제 픽셀 hit, 양쪽 면적 점수와 chord 검증"],
        ["모바일 LAN", "모바일에서 혼자 하기를 눌러도 시작되지 않는 원인을 수정한다.", "비보안 HTTP에서 randomUUID 미지원 시 getRandomValues 기반 UUID v4 폴백"],
    ], [31, 68, 79], header=True, font_size=8.5)
    add_subheading(doc, "사람이 거절한 대안")
    add_bullets(doc, [
        "더블클릭 포획: 이동하는 음식에서 두 번째 명중이 어려워 제외.",
        "포획 대기 중 일시정지: 점수·시간 악용과 리듬 중단 위험으로 제외.",
        "원형 판정: 비원형 음식에서 보이는 그림과 점수 범위가 달라 알파 실루엣으로 교체.",
        "음성의 단일 자동 지표 승인: 발음·강세·문장 연결을 놓쳐 사람 청취와 블라인드 비교로 교체.",
    ])
    add_note(doc, "실기기 QA", "길게 누르기와 이동 임계값은 Playwright 모바일 터치 에뮬레이션으로 검증했다. 실제 iOS·Android 다기기 검증은 제출 전 남은 QA다.", fill=YELLOW)

    doc.add_page_break()
    add_page_title(doc, "03", "AI 음식 이미지 50종")
    add_numbers(doc, [
        "음식 식별성, 먹음직스러움, 비원형 외곽, 작은 화면 가독성을 기준으로 프롬프트를 작성했다.",
        "기존 승인 8종은 유지하고 42종을 신규 생성·교정했다.",
        "녹색·자홍색 크로마키를 로컬에서 제거하고 despill·soft matte를 적용했다.",
        "알파 외곽 기준으로 크롭하고 종횡비를 보존한 채 긴 변 512px로 정규화했다.",
        "투명 WebP, 크기 예산, 카탈로그 50종 일치, 다양한 종횡비를 자동 검사했다.",
        "실제 Phaser 첫 라운드에 메뉴를 하나씩 고정해 표시 크기·알파 마스크·베기/포획 중심을 전수 검수했다.",
    ])
    add_table(doc, [
        ["지표", "결과"],
        ["음식 이미지", "50개"],
        ["기존 승인 / 신규·교정", "8 / 42"],
        ["전체 용량", "2,316,336 B"],
        ["최대 파일", "74,432 B"],
        ["가장 큰 20종", "1,149,052 B"],
        ["타이틀 자산", "5개 / 152,150 B"],
    ], [68, 110], header=True, first_col_fill=LILAC, font_size=9.2)
    add_subheading(doc, "사람의 선별")
    add_body(doc, "짜장면의 과도한 면 리프트, 삼계탕의 집게, 비빔밥의 부자연스러운 젓가락 양, 순대국의 숟가락, 삼겹살의 과도한 비계·넘침 등을 제거하거나 다시 생성했다. 반대로 식별성과 역동성이 좋았던 볶음밥 시안은 유지했다.")
    add_note(doc, "저작물 범위", "제3자의 음식 사진이나 일러스트는 사용하지 않았다.", fill=MINT)

    doc.add_page_break()
    add_page_title(doc, "04", "AI 합성 나레이션 50종")
    add_table(doc, [
        ["항목", "내용"],
        ["서비스", "Microsoft Azure AI Speech Standard S0 · southeastasia"],
        ["보존 음성", "53개 (활성 50개 + historical 3개)"],
        ["활성 모델", "MAI-Voice-2-Flash 45개 / MAI-Voice-2 5개"],
        ["사용 음성", "Haena Flash / Junho Flash / Junho Full"],
    ], [44, 134], header=True, first_col_fill=LILAC, font_size=9.1)
    add_bullets(doc, [
        "합성 전 카탈로그·기존 파일 hash·voice/style·비용 상한을 검사한다.",
        "파일 덮어쓰기 금지, 클립당 1회 요청, retry 0, 키와 문구가 포함된 오류 redaction을 강제한다.",
        "길이만으로 승인하지 않고 발음, 자연스러운 연결, 코미디 강세를 직접 듣는다.",
        "Flash와 Full을 모델명이 보이지 않게 비교하고, 구조가 다른 후보끼리는 별도 결선으로 평가했다.",
        "로컬 trim은 승인된 무음·저에너지 PCM에만 적용하고 retained PCM·hash·구간을 기록했다.",
        "음성 복제, 실제 인물·방송·캐릭터 성대모사는 사용하지 않았다.",
    ])
    add_note(doc, "게임 내 AI 음성 고지", "이 게임의 일부 음식 나레이션은 Microsoft Azure AI Speech로 생성한 AI 합성 음성입니다. 실제 인물의 녹음이나 성대모사가 아닙니다.", fill=YELLOW)

    doc.add_page_break()
    add_page_title(doc, "05", "런타임 분리와 보안")
    add_note(doc, "물리적 분리", "제작 환경의 Codex·ImageGen·Azure 후보 생성과 로컬 후처리는 배포 빌드에 포함되지 않는다. 배포물에는 승인된 정적 WebP 50종과 정적 나레이션 50종만 들어간다.", fill=MINT)
    add_bullets(doc, [
        "게임 실행 중 생성형 AI나 Azure API를 호출하지 않는다.",
        "Azure 키, 엔드포인트, Speech SDK를 배포 번들에 포함하지 않는다.",
        "현재 라운드의 20종 덱에 필요한 이미지와 음원만 선로딩한다.",
        "나레이션은 한 번에 하나만 재생하고 BGM을 -6dB 낮춘 뒤 복구한다.",
        "오디오 load·decode 실패 시 게임을 멈추지 않고 자막만 표시한다.",
        "Firebase 웹 클라이언트 설정은 공개 식별자이며 관리자·서비스 계정 비밀키를 번들에 넣지 않는다.",
    ])
    add_subheading(doc, "멀티플레이 신뢰 경계")
    add_body(doc, "Firestore 규칙은 인증 사용자, 허용 필드, 상태 전이, 방장 권한, 결과 문서의 작성자·불변성과 값 범위를 검증한다. 개인 점수·포획 메뉴는 현재 클라이언트 신뢰이며 서버 리플레이 판정과 App Check는 후속 보강 항목이다.")
    add_subheading(doc, "배포 번들 보안 확인")
    add_table(doc, [
        ["검사", "기대 결과"],
        ["Azure Speech 키·endpoint·SDK", "0건"],
        ["런타임 생성형 AI 호출", "0건"],
        ["승인된 정적 나레이션", "50개"],
        ["historical 음성의 런타임 참조", "0건"],
    ], [75, 103], header=True, first_col_fill=LILAC, font_size=9.2)

    doc.add_page_break()
    add_page_title(doc, "06", "검증 결과와 남은 QA")
    add_table(doc, [
        ["검증", "결과", "범위"],
        ["TypeScript", "제출 commit에서 최종 확인", "현재 제목 변경 범위는 통과했으며 전체 작업 트리는 최종 재실행"],
        ["단위 테스트", "705 pass / 13 skip", "기하·규칙·방·결과·Firebase 코덱·자산·나레이션"],
        ["배포 빌드", "Vite 산출물 생성 통과", "npm run build는 제출 commit에서 타입 검사와 함께 재실행"],
        ["음식 50종", "통과", "카탈로그·알파·용량·덱 선로딩·실제 Phaser 전수 QA"],
        ["Firestore 규칙", "통과", "인증·전이·180초 deadline·늦은 제출 거부·결과 불변"],
        ["나레이션 E2E", "통과", "데스크톱 시나리오 3 pass / 모바일 프로젝트 3 intentional skip"],
    ], [34, 49, 95], header=True, font_size=8.3)
    add_subheading(doc, "제출 전 실기기 QA")
    add_bullets(doc, [
        "실제 3~8대 기기의 네트워크 지연·재접속·공동 결과 동기화",
        "모바일 카메라 권한과 QR 스캔, Firebase App Check 적용",
        "iOS·Android에서 50개 음원의 체감 음량과 BGM duck 편차",
        "기기 시각 차이에 따른 카운트다운 표시와 클라이언트 결과 신뢰 보강",
    ])
    add_note(doc, "최종 제출 전", "테스트 수·빌드 상태는 최종 commit에서 다시 실행한 결과로 갱신합니다.", fill=YELLOW)

    doc.add_page_break()
    add_page_title(doc, "07", "외부 에셋·오픈소스·증빙")
    add_body(doc, "게임 런타임에는 제3자의 음식·배경 이미지, 녹음 음성, 사운드 샘플, 외부 폰트를 사용하지 않았다. 효과음과 128 BPM·64-step 배경음은 외부 저장 음원 없이 Web Audio로 실행 중 생성한다. 제출 문서는 ReportLab(BSD)·pypdf(BSD-3-Clause)·python-docx(MIT)와 Windows 기본 맑은 고딕으로 제작했으며 게임 빌드에는 포함하지 않았다.")
    add_table(doc, [
        ["라이선스", "구성 요소"],
        ["MIT", "Phaser 4.2.1, Vite 8.1.5, Vitest 4.1.10, qrcode 1.5.4, 타입 정의, Firebase CLI, python-docx"],
        ["Apache-2.0", "TypeScript 7.0.2, Playwright 1.62.0, Firebase JavaScript SDK 12.16.0, Rules Unit Testing"],
        ["MIT-CMU", "Pillow 12.3.0"],
        ["서비스 약관", "OpenAI Terms of Use - ImageGen 결과 / Microsoft Azure 계정·public preview 조건 - 합성 음성"],
    ], [42, 136], header=True, font_size=8.8)
    add_subheading(doc, "재현·감사 가능한 근거")
    evidence_table = add_table(doc, [
        ["근거", "위치"],
        ["AI 활용 원장", "docs/source/ai-usage.md"],
        ["에셋·라이선스", "docs/evidence/asset-licenses.md"],
        ["프롬프트 로그", "docs/evidence/ai-prompts/ (25개 작업 기록)"],
        ["음성 생성 이력", "docs/source/narration-generation.md, scripts/narration/"],
        ["소스 저장소", "https://github.com/h200082/oneul-mwo-sseol"],
    ], [48, 130], header=True, first_col_fill=LILAC, font_size=9.0)
    source_cell = evidence_table.cell(5, 1)
    source_cell.text = ""
    source_paragraph = source_cell.paragraphs[0]
    source_paragraph.paragraph_format.space_before = Pt(0)
    source_paragraph.paragraph_format.space_after = Pt(0)
    add_hyperlink(
        source_paragraph,
        "https://github.com/h200082/oneul-mwo-sseol",
        "https://github.com/h200082/oneul-mwo-sseol",
    )
    add_note(doc, "최종본 반영 항목", "GitHub Pages 실제 URL · 최종 commit SHA · 제출일. AI 문서에는 별도 YouTube 입력란을 두지 않고 게임 소개 문서의 영상 링크를 참조합니다.", fill=YELLOW)

    add_final_information_page(doc, include_youtube=False)
    doc.save(AI_DOCX)
    return AI_DOCX


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [build_game_document(), build_ai_document()]
    for path in outputs:
        print(path)


if __name__ == "__main__":
    main()
